#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
文件目录推测系统 (File Directory Predictor)
==========================================
基于大模型（视觉模型 VLM + 语义模型 LLM）根据文件内容推测文件应归属的目录。

支持文件类型:
  - PDF (.pdf)
  - DOCX (.docx)
  - DOC  (.doc)  —— 通过 libreoffice 转换
  - 图片 (.png / .jpg / .jpeg / .gif / .webp / .bmp)
  - 纯文本 (.txt / .csv / .md / .json ...)

工作流程:
  1. CatalogLoader   : 加载编目规则 xlsx，构建可检索的目录索引
  2. ContentExtractor: 按文件类型提取文本 + 渲染图像
  3. VLMAnalyzer     : 视觉模型分析图像/PDF页面，提取视觉特征与 OCR 文本
  4. LLMAnalyzer     : 语义模型分析文本内容，理解文档类型与案件类别
  5. DirectoryPredictor: 融合 VLM + LLM 结果，匹配编目规则，输出推测目录

用法:
  python3 file_directory_predictor.py <文件路径> [--catalog <编目xlsx>] [--json]
  python3 file_directory_predictor.py --batch <目录> [--catalog <编目xlsx>]
  python3 file_directory_predictor.py <文件路径> --output <结果文件> --log <日志文件>

示例:
  python3 file_directory_predictor.py /path/to/起诉书.pdf
  python3 file_directory_predictor.py ./scan.png --catalog ./catalog-mapping.xlsx --json
"""

import argparse
import base64
import contextlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import time
from dataclasses import dataclass, field, asdict
from pathlib import Path
from typing import Optional

import pandas as pd

from smart_case_filing.model_client import LegacyFunctionModelClient
from smart_case_filing.agent.legacy_tools import build_legacy_tool_registry
from smart_case_filing.agent.review import ReviewPackageWriter, build_review_payload
from smart_case_filing.agent.run_manager import AgentRunManager
from smart_case_filing.agent.runner import AgentRunner
from smart_case_filing.agent.retry import RetryPolicy
from smart_case_filing.agent.state import AgentState, AgentTraceStore
from smart_case_filing.agent.preflight import check_model_preflight

try:
    import requests
except ImportError:
    requests = None

# ---------------------------------------------------------------------------
# 全局配置
# ---------------------------------------------------------------------------
PROGRAM_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = Path(os.getenv("SMART_CASE_FILING_ROOT", str(PROGRAM_DIR))).expanduser()
DEFAULT_CATALOG = PROJECT_ROOT / "catalog-mapping.xlsx"
TMP_DIR = Path(os.getenv("SMART_CASE_FILING_TMP_DIR", str(PROJECT_ROOT / "_tmp_predict"))).expanduser()
TMP_DIR.mkdir(parents=True, exist_ok=True)

# VLM 分析时最多渲染的 PDF 页数（控制成本）
MAX_PDF_PAGES_FOR_VLM = 10
# 提取文本送入 LLM 的最大字符数
MAX_TEXT_CHARS = 6000
# CLI 调用超时（秒）
CLI_TIMEOUT = 180
DEFAULT_OUTPUT_FILE = "file_directory_predictor_output.txt"
DEFAULT_LOG_FILE = "file_directory_predictor.log"


# ---------------------------------------------------------------------------
# 工具函数
# ---------------------------------------------------------------------------
def _env(name: str) -> str:
    """读取环境变量并去除首尾空白。"""
    return os.getenv(name, "").strip()


def _get_model_config(kind: str) -> dict:
    """读取 OpenAI-compatible 模型配置。

    kind 为 chat 或 vision。专用环境变量优先，缺省回退到通用 AI_*。
    """
    prefix = "AI_CHAT" if kind == "chat" else "AI_VISION"
    base_url = _env(f"{prefix}_BASE_URL") or _env("AI_BASE_URL")
    api_key = _env(f"{prefix}_API_KEY") or _env("AI_API_KEY")
    model = _env(f"{prefix}_MODEL") or _env("AI_MODEL")
    return {
        "base_url": base_url.rstrip("/"),
        "api_key": api_key,
        "model": model,
    }


def _has_openai_config(config: dict) -> bool:
    return bool(config.get("base_url") and config.get("api_key") and config.get("model"))


def _debug_model_error(message: str):
    if _env("AI_DEBUG") == "1":
        print(f"[AI_DEBUG] {message}", file=sys.stderr)


class _Tee:
    """Write a stream to both the original console stream and a file."""

    def __init__(self, *streams):
        self.streams = streams

    def write(self, data: str):
        for stream in self.streams:
            stream.write(data)
        return len(data)

    def flush(self):
        for stream in self.streams:
            stream.flush()


@contextlib.contextmanager
def _save_cli_streams(output_path: Path, log_path: Path):
    output_path = Path(output_path)
    log_path = Path(log_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    log_path.parent.mkdir(parents=True, exist_ok=True)

    original_stdout = sys.stdout
    original_stderr = sys.stderr
    with output_path.open("w", encoding="utf-8") as output_file, \
            log_path.open("w", encoding="utf-8") as log_file, \
            contextlib.redirect_stdout(_Tee(original_stdout, output_file)), \
            contextlib.redirect_stderr(_Tee(original_stderr, log_file)):
        yield


def _post_openai_compatible(messages: list, config: dict, timeout: int) -> str:
    """调用 OpenAI-compatible /chat/completions，返回 message.content。"""
    if requests is None:
        _debug_model_error("requests is not installed; install with: pip install requests")
        return ""
    url = f"{config['base_url']}/chat/completions"
    headers = {
        "Authorization": f"Bearer {config['api_key']}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": config["model"],
        "messages": messages,
    }
    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=timeout)
        if resp.status_code >= 400:
            _debug_model_error(f"HTTP {resp.status_code}: {resp.text[:500]}")
            return ""
        data = resp.json()
        return data.get("choices", [{}])[0].get("message", {}).get("content", "") or ""
    except requests.RequestException as e:
        _debug_model_error(f"request failed: {e}")
        return ""
    except Exception as e:
        _debug_model_error(f"response parse failed: {e}")
        return ""


def _run_openai_compatible_chat(prompt: str, system: Optional[str],
                                config: dict, timeout: int) -> str:
    messages = []
    if system:
        messages.append({"role": "system", "content": system})
    messages.append({"role": "user", "content": prompt})
    return _post_openai_compatible(messages, config, timeout)


def _run_openai_compatible_vision(prompt: str, image_paths: list,
                                  config: dict, timeout: int) -> str:
    content = [{"type": "text", "text": prompt}]
    for p in image_paths:
        try:
            content.append({
                "type": "image_url",
                "image_url": {"url": _image_to_base64(Path(p))},
            })
        except Exception as e:
            _debug_model_error(f"image encode failed: {p}: {e}")
            return ""
    return _post_openai_compatible([{"role": "user", "content": content}], config, timeout)


def _run_zai_chat(prompt: str, system: Optional[str] = None,
                  thinking: bool = False, timeout: int = CLI_TIMEOUT) -> str:
    """调用文本模型，优先使用 OpenAI-compatible HTTP API，未配置时回退 z-ai CLI。"""
    # 安全清洗：去除 null 字节等会导致 subprocess/HTTP 请求异常的字符
    prompt = (prompt or "").replace("\x00", "")
    system = (system or "").replace("\x00", "") if system else None

    config = _get_model_config("chat")
    if _has_openai_config(config):
        return _run_openai_compatible_chat(prompt, system, config, timeout)

    return _run_zai_chat_cli(prompt, system=system, thinking=thinking, timeout=timeout)


def _run_zai_chat_cli(prompt: str, system: Optional[str] = None,
                      thinking: bool = False, timeout: int = CLI_TIMEOUT) -> str:
    """调用 z-ai chat CLI，返回模型回复文本。

    为避免命令行参数过长（ARG_MAX 限制），当 prompt 较长时写入临时文件，
    再用 shell 的 $(cat file) 方式注入。这里改用更稳妥的 Node.js 内联脚本
    调用 z-ai-web-dev-sdk，直接读取文件内容作为 prompt。
    """
    out_file = TMP_DIR / f"chat_{int(time.time() * 1000)}.json"
    # 短 prompt 直接用命令行参数；长 prompt 写入文件后用 cat 注入
    if len(prompt) < 3000 and (not system or len(system) < 3000):
        cmd = ["z-ai", "chat", "-p", prompt, "-o", str(out_file)]
        if system:
            cmd += ["-s", system]
        if thinking:
            cmd += ["-t"]
        try:
            subprocess.run(cmd, capture_output=True, text=True, timeout=timeout, check=False)
        except subprocess.TimeoutExpired:
            return ""
    else:
        # 长 prompt：写入临时文件，用 bash -c 'z-ai chat -p "$(cat file)" ...'
        prompt_file = TMP_DIR / f"prompt_{int(time.time() * 1000)}.txt"
        system_file = None
        try:
            prompt_file.write_text(prompt, encoding="utf-8")
            shell_cmd = f'z-ai chat -p "$(cat {prompt_file})" -o {out_file}'
            if system:
                system_file = TMP_DIR / f"system_{int(time.time() * 1000)}.txt"
                system_file.write_text(system, encoding="utf-8")
                shell_cmd += f' -s "$(cat {system_file})"'
            if thinking:
                shell_cmd += " -t"
            try:
                subprocess.run(["bash", "-c", shell_cmd], capture_output=True,
                               text=True, timeout=timeout, check=False)
            except subprocess.TimeoutExpired:
                return ""
        finally:
            prompt_file.unlink(missing_ok=True)
            if system_file:
                system_file.unlink(missing_ok=True)
    if not out_file.exists():
        return ""
    try:
        data = json.loads(out_file.read_text(encoding="utf-8"))
        return data.get("choices", [{}])[0].get("message", {}).get("content", "") or ""
    except Exception:
        return ""
    finally:
        try:
            out_file.unlink(missing_ok=True)
        except Exception:
            pass


def _run_zai_vision(prompt: str, image_paths: list, thinking: bool = False,
                    timeout: int = CLI_TIMEOUT) -> str:
    """调用视觉模型，优先使用 OpenAI-compatible HTTP API，未配置时回退 z-ai CLI。"""
    prompt = (prompt or "").replace("\x00", "")
    config = _get_model_config("vision")
    if _has_openai_config(config):
        return _run_openai_compatible_vision(prompt, image_paths, config, timeout)

    return _run_zai_vision_cli(prompt, image_paths, thinking=thinking, timeout=timeout)


def _run_zai_vision_cli(prompt: str, image_paths: list, thinking: bool = False,
                        timeout: int = CLI_TIMEOUT) -> str:
    """调用 z-ai vision CLI 分析一张或多张图片，返回模型回复文本。"""
    if not image_paths:
        return ""
    out_file = TMP_DIR / f"vision_{int(time.time() * 1000)}.json"
    cmd = ["z-ai", "vision", "-p", prompt, "-o", str(out_file)]
    for p in image_paths:
        cmd += ["-i", str(p)]
    if thinking:
        cmd += ["-t"]
    try:
        subprocess.run(cmd, capture_output=True, text=True, timeout=timeout, check=False)
    except subprocess.TimeoutExpired:
        return ""
    if not out_file.exists():
        return ""
    try:
        data = json.loads(out_file.read_text(encoding="utf-8"))
        return data.get("choices", [{}])[0].get("message", {}).get("content", "") or ""
    except Exception:
        return ""
    finally:
        try:
            out_file.unlink(missing_ok=True)
        except Exception:
            pass


def _default_model_client():
    return LegacyFunctionModelClient(
        lambda prompt, system=None, thinking=False, timeout=CLI_TIMEOUT: _run_zai_chat(
            prompt, system=system, thinking=thinking, timeout=timeout
        ),
        lambda prompt, image_paths, thinking=False, timeout=CLI_TIMEOUT: _run_zai_vision(
            prompt, image_paths, thinking=thinking, timeout=timeout
        ),
    )


def _image_to_base64(path: Path) -> str:
    """读取图片并转为 data URI。"""
    mime = {
        ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".gif": "image/gif", ".webp": "image/webp", ".bmp": "image/bmp",
    }.get(path.suffix.lower(), "image/png")
    b64 = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime};base64,{b64}"


# ---------------------------------------------------------------------------
# 1. 编目规则加载器
# ---------------------------------------------------------------------------
@dataclass
class CatalogEntry:
    """单条编目规则。"""
    case_type: str               # 案件类型，如 "刑事一审案件编目规则"
    volume: str                  # 卷宗，"正卷" / "副卷"
    second_level_directory: str  # 二级目录
    constraint: str              # 约束，如 "必选"
    material_category: str       # 材料类别
    catalog_name_example: str    # 编目名称示例


@dataclass
class CatalogIndex:
    """编目规则索引，便于快速检索。"""
    entries: list  # list[CatalogEntry]
    # 按材料类别聚合：material_category -> list[(case_type, volume, second_level_directory, example)]
    by_category: dict = field(default_factory=dict)
    # 按案件类型聚合：case_type -> list[CatalogEntry]
    by_case_type: dict = field(default_factory=dict)
    # 所有案件类型
    case_types: list = field(default_factory=list)
    # 所有材料类别（去重）
    categories: list = field(default_factory=list)

    def build(self):
        self.by_category.clear()
        self.by_case_type.clear()
        for e in self.entries:
            key = e.material_category or ""
            self.by_category.setdefault(key, []).append(e)
            self.by_case_type.setdefault(e.case_type, []).append(e)
        self.case_types = sorted({e.case_type for e in self.entries})
        self.categories = sorted({e.material_category for e in self.entries if e.material_category})
        return self

    def summary_for_llm(self) -> str:
        """生成给 LLM 用的编目摘要（控制长度）。"""
        lines = []
        lines.append(f"# 编目规则摘要（共 {len(self.entries)} 条，{len(self.case_types)} 种案件类型）")
        lines.append("## 案件类型列表")
        for ct in self.case_types:
            lines.append(f"- {ct}")
        lines.append("")
        lines.append("## 材料类别 -> 对应案件类型/二级目录")
        for cat in self.categories:
            rows = self.by_category.get(cat, [])
            case_set = sorted({r.case_type for r in rows})
            dir_set = sorted({r.second_level_directory for r in rows})
            vol_set = sorted({r.volume for r in rows})
            ex = next((r.catalog_name_example for r in rows if r.catalog_name_example), "")
            lines.append(f"- 【{cat}】 卷宗:{'/'.join(vol_set)} | 二级目录:{' / '.join(dir_set[:3])} | 适用案件:{' / '.join(case_set)} | 示例:{ex}")
        return "\n".join(lines)

    def search_candidates(self, keywords: list, top_n: int = 25,
                          case_number: str = "") -> list:
        """根据关键词列表，本地模糊检索最相关的编目条目（两阶段检索的第一阶段）。

        若能从案号（通常来自父目录名）识别案件类型，则直接返回该案件类型下
        的全部编目规则，不再按 top_n 截断。

        打分策略：对每条编目条目，统计其各字段中命中的关键词数，取 top_n。
        keywords 来自 LLM/VLM 提取的文档画像（文书类型、案件线索、关键短语等）。
        """
        case_type = self._case_type_from_case_number(case_number)
        if case_type:
            entries = self.by_case_type.get(case_type, [])
            if entries:
                return list(entries)

        if not keywords:
            # 无关键词时返回各案件类型的代表性条目（封面/目录等通用项）
            seen = set()
            reps = []
            for e in self.entries:
                key = (e.case_type, e.material_category)
                if key not in seen and ("封面" in e.material_category or "目录" in e.material_category):
                    seen.add(key)
                    reps.append(e)
            return reps[:top_n]
        scored = []
        kw_list = [k.strip() for k in keywords if k and k.strip()]
        for e in self.entries:
            hay = " ".join([e.case_type, e.volume, e.second_level_directory,
                            e.material_category, e.catalog_name_example]).lower()
            score = 0
            for kw in kw_list:
                kwl = kw.lower()
                if not kwl:
                    continue
                if kwl in hay:
                    score += 2
                    # 材料类别命中加权
                    if kwl in e.material_category.lower():
                        score += 3
                    # 编目示例命中加权
                    if kwl in e.catalog_name_example.lower():
                        score += 1
            if score > 0:
                scored.append((score, e))
        scored.sort(key=lambda x: -x[0])
        # 去重：同 (case_type, material_category, second_level_directory) 只保留得分最高的
        seen = set()
        result = []
        for score, e in scored:
            key = (e.case_type, e.material_category, e.second_level_directory, e.volume)
            if key not in seen:
                seen.add(key)
                result.append(e)
            if len(result) >= top_n:
                break
        return result

    def _case_type_from_case_number(self, case_number: str) -> str:
        """从案号文本中推导编目表中的案件类型名称。"""
        text = str(case_number or "").strip()
        if not text:
            return ""

        mappings = [
            (r"民初", "民事一审案件编目规则"),
            (r"民终", "民事二审案件编目规则"),
            (r"民申", "民事申请再审审查案件编目规则"),
            (r"民再", "民事再审案件编目规则"),
            (r"刑初", "刑事一审案件编目规则"),
            (r"刑终", "刑事二审案件编目规则"),
            (r"刑申", "刑事申诉再审审查案件编目规则"),
            (r"刑再", "刑事再审案件编目规则"),
            (r"行初", "行政一审案件编目规则"),
            (r"行终", "行政二审案件编目规则"),
            (r"行申", "行政申请再审审查案件编目规则"),
            (r"行再", "行政再审案件编目规则"),
            (r"执", "首次执行案件编目规则"),
        ]
        for pattern, case_type in mappings:
            if re.search(pattern, text):
                return case_type
        return ""


class CatalogLoader:
    """从 xlsx 加载编目规则。"""

    def __init__(self, xlsx_path: Path):
        self.xlsx_path = Path(xlsx_path)

    def load(self) -> CatalogIndex:
        df = pd.read_excel(self.xlsx_path)
        df = df.fillna("")
        entries = []
        for _, row in df.iterrows():
            entries.append(CatalogEntry(
                case_type=str(row.get("case_type", "")).strip(),
                volume=str(row.get("volume", "")).strip(),
                second_level_directory=str(row.get("second_level_directory", "")).strip(),
                constraint=str(row.get("constraint", "")).strip(),
                material_category=str(row.get("material_category", "")).strip(),
                catalog_name_example=str(row.get("catalog_name_example", "")).strip(),
            ))
        idx = CatalogIndex(entries=entries).build()
        return idx


# ---------------------------------------------------------------------------
# 2. 文件内容提取器
# ---------------------------------------------------------------------------
@dataclass
class FileContent:
    """提取出的文件内容。"""
    file_path: Path
    file_type: str            # pdf / docx / doc / image / text / unknown
    text: str = ""            # 提取的文本
    image_paths: list = field(default_factory=list)  # 渲染出的图片路径（供 VLM）
    page_count: int = 0
    extract_error: str = ""

    def has_visual(self) -> bool:
        return len(self.image_paths) > 0

    def text_preview(self, n: int = 1500) -> str:
        # 清除 null 字节等控制字符，避免传给 subprocess 时报错
        clean = self.text.replace("\x00", "").replace("\r", "")
        # 去除其他不可见控制字符（保留换行和制表符）
        clean = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f]", "", clean)
        return clean[:n]


class ContentExtractor:
    """按文件类型提取文本与图像。"""

    IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"}
    TEXT_EXTS = {".txt", ".csv", ".tsv", ".md", ".json", ".log", ".xml", ".html", ".htm"}
    SHEET_EXTS = {".xlsx", ".xlsm", ".xls"}

    def extract(self, file_path: Path) -> FileContent:
        file_path = Path(file_path)
        ext = file_path.suffix.lower()
        if ext == ".pdf":
            return self._extract_pdf(file_path)
        if ext == ".docx":
            return self._extract_docx(file_path)
        if ext == ".doc":
            return self._extract_doc(file_path)
        if ext in self.IMAGE_EXTS:
            return self._extract_image(file_path)
        if ext in self.SHEET_EXTS:
            return self._extract_sheet(file_path)
        if ext in self.TEXT_EXTS:
            return self._extract_text_file(file_path)
        # 兜底：尝试当文本读
        return self._extract_text_file(file_path)

    # ---- PDF ----
    def _extract_pdf(self, path: Path) -> FileContent:
        fc = FileContent(file_path=path, file_type="pdf")
        # 文本提取：优先 PyMuPDF（快且稳）
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(path))
            fc.page_count = doc.page_count
            parts = []
            for i, page in enumerate(doc):
                parts.append(page.get_text("text"))
                if i >= 30:  # 限制页数避免超大文件
                    break
            fc.text = "\n".join(parts).strip()
            doc.close()
        except Exception as e:
            fc.extract_error = f"PyMuPDF extract failed: {e}"

        # 渲染前 N 页为图片供 VLM
        try:
            import fitz
            doc = fitz.open(str(path))
            n_pages = min(MAX_PDF_PAGES_FOR_VLM, doc.page_count)
            for i in range(n_pages):
                page = doc[i]
                # 适当放大以提升 OCR 质量
                mat = fitz.Matrix(2.0, 2.0)
                pix = page.get_pixmap(matrix=mat)
                img_path = TMP_DIR / f"{path.stem}_page{i + 1}.png"
                pix.save(str(img_path))
                fc.image_paths.append(img_path)
            doc.close()
        except Exception as e:
            if not fc.extract_error:
                fc.extract_error = f"PDF render failed: {e}"
        return fc

    # ---- DOCX ----
    def _extract_docx(self, path: Path) -> FileContent:
        fc = FileContent(file_path=path, file_type="docx")
        try:
            import docx
            d = docx.Document(str(path))
            parts = [p.text for p in d.paragraphs if p.text.strip()]
            # 表格内容
            for tbl in d.tables:
                for row in tbl.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            fc.text = "\n".join(parts).strip()
            fc.page_count = 1
        except Exception as e:
            fc.extract_error = f"docx extract failed: {e}"
        return fc

    # ---- DOC (老格式，先转 docx) ----
    def _extract_doc(self, path: Path) -> FileContent:
        fc = FileContent(file_path=path, file_type="doc")
        try:
            out_dir = TMP_DIR / f"docconv_{path.stem}_{int(time.time())}"
            out_dir.mkdir(parents=True, exist_ok=True)
            r = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "docx",
                 "--outdir", str(out_dir), str(path)],
                capture_output=True, text=True, timeout=120, check=False,
            )
            converted = out_dir / (path.stem + ".docx")
            if converted.exists():
                inner = self._extract_docx(converted)
                fc.text = inner.text
                fc.extract_error = inner.extract_error
                fc.image_paths = inner.image_paths
                fc.page_count = inner.page_count
                # 清理临时
                shutil.rmtree(out_dir, ignore_errors=True)
            else:
                fc.extract_error = f"libreoffice convert failed: {r.stderr[:200]}"
        except Exception as e:
            fc.extract_error = f"doc convert failed: {e}"
        return fc

    # ---- 图片 ----
    def _extract_image(self, path: Path) -> FileContent:
        fc = FileContent(file_path=path, file_type="image", image_paths=[path], page_count=1)
        # 不做本地 OCR，交给 VLM
        return fc

    # ---- 电子表格 (xlsx/xls) ----
    def _extract_sheet(self, path: Path) -> FileContent:
        fc = FileContent(file_path=path, file_type="spreadsheet", page_count=1)
        try:
            df = pd.read_excel(path, sheet_name=None, nrows=200)  # 所有 sheet，限 200 行
            parts = []
            for sheet_name, frame in df.items():
                parts.append(f"[工作表: {sheet_name}]")
                parts.append(frame.fillna("").astype(str).to_csv(index=False, sep="|"))
            fc.text = "\n".join(parts).strip()
        except Exception as e:
            fc.extract_error = f"xlsx read failed: {e}"
        return fc

    # ---- 纯文本 ----
    def _extract_text_file(self, path: Path) -> FileContent:
        fc = FileContent(file_path=path, file_type="text", page_count=1)
        try:
            for enc in ("utf-8", "gbk", "gb18030", "latin-1"):
                try:
                    fc.text = path.read_text(encoding=enc).strip()
                    break
                except UnicodeDecodeError:
                    continue
        except Exception as e:
            fc.extract_error = f"text read failed: {e}"
        return fc


# ---------------------------------------------------------------------------
# 3. VLM 视觉分析器
# ---------------------------------------------------------------------------
class VLMAnalyzer:
    """使用视觉模型分析图像/PDF页面，提取视觉特征与 OCR 文本。"""

    SYSTEM = (
        "你是法院卷宗材料视觉分析专家。你擅长识别法律文书的版式、印章、签名、表格、"
        "手写体、打印体等视觉特征，并能准确 OCR 出图片中的中文文字。"
    )

    def __init__(self, model_client=None):
        self.model_client = model_client or _default_model_client()

    def analyze(self, fc: FileContent) -> dict:
        """返回 {visual_description, ocr_text, doc_type_guess, volume_guess, case_clues}。"""
        if not fc.has_visual():
            return {"available": False, "reason": "无可分析图像"}
        # 限制图片数量，避免 token 爆炸
        imgs = fc.image_paths[:MAX_PDF_PAGES_FOR_VLM]
        prompt = self._build_prompt(fc)
        raw = self.model_client.vision(prompt, imgs, thinking=False, timeout=CLI_TIMEOUT)
        return self._parse(raw)

    def _build_prompt(self, fc: FileContent) -> str:
        return (
            "请分析这张/这几张法院卷宗材料图片，按如下 JSON 格式严格输出（不要输出 JSON 以外的内容）：\n"
            "{\n"
            '  "doc_type_guess": "推测的文书类型，如 卷宗封面/卷内目录/起诉书/判决书/裁定书/调解书/'
            '立案登记表/案件审判流程管理信息表/送达回证/笔录/鉴定意见/证据材料/委托手续 等",\n'
            '  "volume_guess": "正卷 或 副卷 或 未知（依据是否涉及内部审批、合议笔录等判断）",\n'
            '  "case_clues": "案件类型线索，如 刑事/民事/行政/执行/再审 等关键词",\n'
            '  "visual_features": "视觉特征描述，如 红色公章/手写签名/表格/印刷体/骑缝章 等",\n'
            '  "ocr_text": "图片中识别到的关键文字内容（前800字即可）",\n'
            '  "confidence": "high/medium/low"\n'
            "}\n"
            "注意：只输出 JSON，不要有任何解释性文字。"
        )

    def _parse(self, raw: str) -> dict:
        if not raw:
            return {"available": False, "reason": "VLM 无返回"}
        # 尝试提取 JSON
        raw = raw.strip()
        # 去掉可能的 ```json 包裹
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
        try:
            obj = json.loads(raw)
            obj["available"] = True
            return obj
        except Exception:
            return {"available": True, "raw": raw, "parse_error": True}


# ---------------------------------------------------------------------------
# 4. LLM 语义分析器
# ---------------------------------------------------------------------------
class LLMAnalyzer:
    """使用语义模型分析文本内容，理解文档类型与案件类别。"""

    SYSTEM = (
        "你是法院卷宗编目专家，精通各类案件（刑事/民事/行政/执行/再审）的卷宗材料分类规则。"
        "你能根据文书内容准确判断其所属的案件类型、卷宗（正卷/副卷）、二级目录与材料类别。"
    )

    def __init__(self, model_client=None):
        self.model_client = model_client or _default_model_client()

    def analyze_text(self, fc: FileContent) -> dict:
        """分析提取出的文本，返回文档语义画像。"""
        if not fc.text or not fc.text.strip():
            return {"available": False, "reason": "无文本可分析"}
        text = fc.text[:MAX_TEXT_CHARS]
        prompt = (
            "请分析以下从文件中提取的文本内容，按如下 JSON 格式严格输出（不要输出 JSON 以外的内容）：\n"
            "{\n"
            '  "doc_type_guess": "推测的文书类型，如 起诉书/判决书/裁定书/调解书/立案登记表/'
            '案件审判流程管理信息表/送达回证/庭审笔录/合议庭评议笔录/鉴定意见/委托书/证据 等",\n'
            '  "volume_guess": "正卷 或 副卷 或 未知",\n'
            '  "case_clues": "案件类型线索关键词，如 刑事/民事/行政/执行/再审/一审/二审 等",\n'
            '  "key_phrases": "内容中能体现文书性质的关键短语，逗号分隔",\n'
            '  "summary": "一句话概括该文件内容",\n'
            '  "confidence": "high/medium/low"\n'
            "}\n"
            f"文件名: {fc.file_path.name}\n"
            f"文件类型: {fc.file_type}\n"
            f"文本内容:\n{text}\n"
            "注意：只输出 JSON，不要有任何解释性文字。"
        )
        raw = self.model_client.chat(prompt, system=self.SYSTEM, thinking=False, timeout=CLI_TIMEOUT)
        return self._parse(raw)

    def _parse(self, raw: str) -> dict:
        if not raw:
            return {"available": False, "reason": "LLM 无返回"}
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
        try:
            obj = json.loads(raw)
            obj["available"] = True
            return obj
        except Exception:
            return {"available": True, "raw": raw, "parse_error": True}


# ---------------------------------------------------------------------------
# 5. 目录推测器（融合 VLM + LLM）
# ---------------------------------------------------------------------------
@dataclass
class PredictionResult:
    """推测结果。"""
    file_path: str
    file_type: str
    predicted_case_type: str = ""
    predicted_volume: str = ""
    predicted_second_level_directory: str = ""
    predicted_material_category: str = ""
    predicted_catalog_example: str = ""
    confidence: str = "low"
    reasoning: str = ""
    vlm_analysis: dict = field(default_factory=dict)
    llm_analysis: dict = field(default_factory=dict)
    matched_entries: list = field(default_factory=list)

    def to_dict(self) -> dict:
        return asdict(self)

    def predicted_path(self) -> str:
        parts = [p for p in [
            self.predicted_case_type,
            self.predicted_volume,
            self.predicted_second_level_directory,
            self.predicted_material_category,
        ] if p]
        return " / ".join(parts)


class DirectoryPredictor:
    """融合 VLM 与 LLM 分析结果，匹配编目规则，输出推测目录。"""

    def __init__(self, catalog: CatalogIndex, model_client=None):
        self.catalog = catalog
        self.model_client = model_client or _default_model_client()
        self.vlm = VLMAnalyzer(self.model_client)
        self.llm = LLMAnalyzer(self.model_client)

    def predict(self, file_path: Path) -> PredictionResult:
        file_path = Path(file_path)
        extractor = ContentExtractor()
        fc = extractor.extract(file_path)
        result = PredictionResult(file_path=str(file_path), file_type=fc.file_type)

        # 并行思路：先 VLM 后 LLM（这里串行调用，简单可靠）
        vlm_res = self.vlm.analyze(fc)
        llm_res = self.llm.analyze_text(fc)
        result.vlm_analysis = vlm_res
        result.llm_analysis = llm_res

        # 融合 + 匹配编目
        fusion = self._fuse(vlm_res, llm_res, fc)
        match = self._match_catalog(fusion)
        result.predicted_case_type = match.get("case_type", "")
        result.predicted_volume = match.get("volume", "")
        result.predicted_second_level_directory = match.get("second_level_directory", "")
        result.predicted_material_category = match.get("material_category", "")
        result.predicted_catalog_example = match.get("catalog_name_example", "")
        result.confidence = match.get("confidence", "low")
        result.reasoning = match.get("reasoning", "")
        result.matched_entries = match.get("matched_entries", [])
        return result

    # ---- 融合 VLM + LLM ----
    def _fuse(self, vlm_res: dict, llm_res: dict, fc: FileContent) -> dict:
        """把两路分析结果合并成统一的文档画像。"""
        doc_type = ""
        volume = ""
        case_clues = ""
        key_info = ""
        confidence = "low"

        if vlm_res.get("available") and not vlm_res.get("parse_error"):
            doc_type = vlm_res.get("doc_type_guess", "") or doc_type
            volume = vlm_res.get("volume_guess", "") or volume
            case_clues = vlm_res.get("case_clues", "") or case_clues
            key_info += "视觉特征: " + (vlm_res.get("visual_features", "") or "") + "\n"
            key_info += "OCR文本: " + (vlm_res.get("ocr_text", "") or "") + "\n"
            if vlm_res.get("confidence") == "high":
                confidence = "medium"

        if llm_res.get("available") and not llm_res.get("parse_error"):
            doc_type = llm_res.get("doc_type_guess", "") or doc_type
            volume = llm_res.get("volume_guess", "") or volume
            case_clues = (case_clues + " " + (llm_res.get("case_clues", "") or "")).strip()
            key_info += "关键短语: " + (llm_res.get("key_phrases", "") or "") + "\n"
            key_info += "内容摘要: " + (llm_res.get("summary", "") or "") + "\n"
            if llm_res.get("confidence") == "high":
                confidence = "high" if confidence == "medium" else "medium"

        return {
            "doc_type": doc_type,
            "volume": volume,
            "case_clues": case_clues,
            "key_info": key_info,
            "case_number": fc.file_path.parent.name,
            "file_name": fc.file_path.name,
            "file_type": fc.file_type,
            "text_preview": fc.text_preview(800),
            "confidence": confidence,
        }

    # ---- 匹配编目规则（两阶段检索）----
    def _match_catalog(self, fusion: dict) -> dict:
        """两阶段检索匹配编目规则。

        阶段1（本地检索）：优先根据案号推导案件类型，加载该案件类型下所有编目规则；
        若无法推导，再从 fusion 画像中提取关键词进行本地模糊检索。
        阶段2（LLM 精选）：把文件画像 + 候选条目交给 LLM，让其选出最匹配的一条。
        这样避免把全部编目摘要塞进 prompt（会超 ARG_MAX）。
        """
        # ---- 阶段1：提取关键词并本地检索 ----
        keywords = self._extract_keywords(fusion)
        candidates = self.catalog.search_candidates(
            keywords,
            case_number=fusion.get("case_number", ""),
        )
        if not candidates:
            # 关键词没命中任何条目，退回到通用条目
            candidates = self.catalog.search_candidates([])

        # 构造候选清单文本（编号 + 字段）
        cand_lines = []
        for i, e in enumerate(candidates, 1):
            cand_lines.append(
                f"[{i}] 案件类型:{e.case_type} | 卷宗:{e.volume} | "
                f"二级目录:{e.second_level_directory} | 材料类别:{e.material_category} | "
                f"示例:{e.catalog_name_example}"
            )
        candidate_text = "\n".join(cand_lines)

        # ---- 阶段2：LLM 从候选中精选 ----
        prompt = (
            "你是法院卷宗编目匹配专家。下面给出一份文件的【分析画像】与【候选编目条目】。\n"
            "请根据文件画像，从候选条目中选出【最匹配】的一条（只能从候选编号中选），并给出理由。\n\n"
            "严格按如下 JSON 格式输出（不要输出 JSON 以外的任何内容）：\n"
            "{\n"
            '  "selected_index": "候选编号(数字)",\n'
            '  "case_type": "对应的案件类型",\n'
            '  "volume": "正卷 或 副卷",\n'
            '  "second_level_directory": "对应的二级目录名称",\n'
            '  "material_category": "对应的材料类别",\n'
            '  "catalog_name_example": "对应的编目名称示例",\n'
            '  "confidence": "high/medium/low",\n'
            '  "reasoning": "一句话说明为什么这样匹配"\n'
            "}\n\n"
            "=== 文件分析画像 ===\n"
            f"文件名: {fusion.get('file_name')}\n"
            f"案号/目录名: {fusion.get('case_number')}\n"
            f"文件类型: {fusion.get('file_type')}\n"
            f"推测文书类型: {fusion.get('doc_type')}\n"
            f"推测卷宗: {fusion.get('volume')}\n"
            f"案件线索: {fusion.get('case_clues')}\n"
            f"画像置信度: {fusion.get('confidence')}\n"
            f"关键信息:\n{fusion.get('key_info')}\n"
            f"文本预览:\n{fusion.get('text_preview')}\n\n"
            f"=== 候选编目条目（共 {len(candidates)} 条）===\n"
            f"{candidate_text}\n\n"
            "匹配要求：\n"
            "1. selected_index 必须是上面候选列表中的编号；\n"
            "2. 其余字段须与该编号对应的条目完全一致；\n"
            "3. 若所有候选都不匹配，confidence 设为 low 并在 reasoning 中说明。\n"
            "注意：只输出 JSON。"
        )
        raw = self.model_client.chat(prompt, system=LLMAnalyzer.SYSTEM, thinking=True, timeout=CLI_TIMEOUT)
        return self._parse_match(raw, candidates)

    def _extract_keywords(self, fusion: dict) -> list:
        """从融合画像中提取用于本地检索的关键词。"""
        kws = []
        # 文书类型（如 起诉书/判决书/卷宗封面/登记表/笔录/裁定书 等）
        dt = fusion.get("doc_type", "")
        if dt:
            kws += re.split(r"[、,，/；;\s]+", dt)
        # 案件线索（如 刑事/民事/行政/一审/二审/再审/执行 等）
        cl = fusion.get("case_clues", "")
        if cl:
            kws += re.split(r"[、,，/；;\s]+", cl)
        # 关键信息中的短语
        ki = fusion.get("key_info", "")
        if ki:
            # 提取中文短语（2-6字）
            kws += re.findall(r"[\u4e00-\u9fa5]{2,6}", ki)
        # 文件名中的线索
        fn = fusion.get("file_name", "")
        if fn:
            kws += re.findall(r"[\u4e00-\u9fa5]{2,8}", fn)
        # 去重、去空、去太短的
        seen = set()
        result = []
        for k in kws:
            k = k.strip()
            if k and len(k) >= 2 and k not in seen:
                seen.add(k)
                result.append(k)
        return result[:40]  # 限制关键词数量

    def _parse_match(self, raw: str, candidates: list = None) -> dict:
        """解析 LLM 的匹配结果，并用候选列表校验 selected_index。"""
        if not raw:
            return {"confidence": "low", "reasoning": "LLM 匹配无返回"}
        raw = raw.strip()
        if raw.startswith("```"):
            raw = re.sub(r"^```(?:json)?\s*", "", raw)
            raw = re.sub(r"\s*```$", "", raw)
        # 尝试提取第一个 JSON 对象
        m = re.search(r"\{[\s\S]*\}", raw)
        if m:
            raw = m.group(0)
        try:
            obj = json.loads(raw)
        except Exception:
            return {"confidence": "low", "reasoning": "LLM 匹配返回解析失败", "raw": raw[:500]}

        # 优先用 selected_index 定位候选条目（最可靠）
        if candidates:
            idx_raw = obj.get("selected_index", "")
            idx = None
            try:
                idx = int(re.search(r"\d+", str(idx_raw)).group())
            except Exception:
                pass
            if idx and 1 <= idx <= len(candidates):
                e = candidates[idx - 1]
                obj["case_type"] = e.case_type
                obj["volume"] = e.volume
                obj["second_level_directory"] = e.second_level_directory
                obj["material_category"] = e.material_category
                obj["catalog_name_example"] = e.catalog_name_example
                obj["matched_entries"] = [asdict(e)]
                return obj

        # 回退：用 LLM 给的字段在候选中匹配
        ct = obj.get("case_type", "")
        vol = obj.get("volume", "")
        sld = obj.get("second_level_directory", "")
        mc = obj.get("material_category", "")
        if candidates:
            matched = [
                asdict(e) for e in candidates
                if e.case_type == ct and e.volume == vol
                and (e.second_level_directory == sld or not sld)
                and (e.material_category == mc or not mc)
            ]
            if not matched:
                matched = [asdict(e) for e in candidates
                           if e.case_type == ct and e.material_category == mc]
        else:
            matched = [
                asdict(e) for e in self.catalog.entries
                if e.case_type == ct and e.volume == vol
                and (e.second_level_directory == sld or not sld)
                and (e.material_category == mc or not mc)
            ]
            if not matched:
                matched = [asdict(e) for e in self.catalog.entries
                           if e.case_type == ct and e.material_category == mc]
        if matched:
            best = matched[0]
            obj["case_type"] = best["case_type"]
            obj["volume"] = best["volume"]
            obj["second_level_directory"] = best["second_level_directory"]
            obj["material_category"] = best["material_category"]
            obj["catalog_name_example"] = best["catalog_name_example"]
        obj["matched_entries"] = matched[:5]
        return obj


# ---------------------------------------------------------------------------
# CLI 入口
# ---------------------------------------------------------------------------
def _print_result(result: PredictionResult, as_json: bool = False):
    if as_json:
        print(json.dumps(result.to_dict(), ensure_ascii=False, indent=2))
        return
    print("=" * 72)
    print(f"文件: {result.file_path}")
    print(f"类型: {result.file_type}")
    print("-" * 72)
    print(f"推测目录: {result.predicted_path() or '(未能推测)'}")
    print(f"编目示例: {result.predicted_catalog_example or '-'}")
    print(f"置信度  : {result.confidence}")
    print(f"理由    : {result.reasoning or '-'}")
    print("-" * 72)
    if result.vlm_analysis.get("available"):
        print("[VLM 视觉分析]")
        for k in ("doc_type_guess", "volume_guess", "case_clues", "visual_features", "confidence"):
            v = result.vlm_analysis.get(k)
            if v:
                print(f"  {k}: {v}")
        ocr = result.vlm_analysis.get("ocr_text", "")
        if ocr:
            print(f"  ocr_text: {ocr[:200]}...")
    else:
        print("[VLM 视觉分析] 不可用 -", result.vlm_analysis.get("reason", ""))
    print("-" * 72)
    if result.llm_analysis.get("available"):
        print("[LLM 语义分析]")
        for k in ("doc_type_guess", "volume_guess", "case_clues", "key_phrases", "summary", "confidence"):
            v = result.llm_analysis.get(k)
            if v:
                print(f"  {k}: {v}")
    else:
        print("[LLM 语义分析] 不可用 -", result.llm_analysis.get("reason", ""))
    print("=" * 72)


def _run_agent_cli(args):
    if getattr(args, "review_decision", ""):
        _run_review_decision_cli(args)
        return

    if getattr(args, "agent_preflight", False):
        print(json.dumps(check_model_preflight(), ensure_ascii=False, indent=2))
        return

    trace_path = Path(args.trace) if args.trace else PROGRAM_DIR / "agent_trace.jsonl"
    if args.resume:
        _run_agent_resume(args)
        return

    if args.batch:
        _run_agent_batch_cli(args)
        return

    catalog_path = Path(args.catalog)
    file_path = Path(args.file) if args.file else None

    if not file_path or not file_path.exists():
        result = {
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "error": f"file does not exist: {file_path}" if file_path else "file is required",
            "trace": str(trace_path),
            "review_output": args.review_output or "",
            "resume": args.resume or "",
        }
        if args.review_output:
            ReviewPackageWriter(Path(args.review_output)).write(build_review_payload(result, str(trace_path)))
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return

    if not catalog_path.exists():
        result = {
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "error": f"catalog does not exist: {catalog_path}",
            "trace": str(trace_path),
            "review_output": args.review_output or "",
            "resume": args.resume or "",
        }
        if args.review_output:
            ReviewPackageWriter(Path(args.review_output)).write(build_review_payload(result, str(trace_path)))
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return

    catalog = CatalogLoader(catalog_path).load()
    trace_store = AgentTraceStore(trace_path)
    registry = build_legacy_tool_registry(catalog)
    runner = AgentRunner(registry, trace_store, retry_policy=_agent_retry_policy_from_args(args))
    run_id = f"agent-{int(time.time() * 1000)}"
    result = runner.run(run_id=run_id, file_path=str(file_path))
    output = _agent_output_from_result(result, trace_path, args.review_output or "", args.resume or "")

    if output["agent_state"] in {AgentState.NEEDS_REVIEW.value, AgentState.FAILED.value} and args.review_output:
        ReviewPackageWriter(Path(args.review_output)).write(build_review_payload(output, str(trace_path)))

    print(json.dumps(output, ensure_ascii=False, indent=2))


def _agent_output_from_result(result: dict, trace_path: Path, review_output: str = "", resume: str = "") -> dict:
    state = result.get("state", AgentState.FAILED)
    state_value = state.value if isinstance(state, AgentState) else str(state)
    output = dict(result.get("prediction") or {})
    output.update({
        "agent_state": state_value,
        "state": state_value,
        "trace": str(trace_path),
        "review_output": review_output,
        "resume": resume,
    })
    if result.get("error"):
        output["error"] = result["error"]
    return output


def _agent_retry_policy_from_args(args) -> RetryPolicy:
    errors = getattr(args, "agent_retry_errors", "") or ""
    retryable_errors = tuple(part.strip().lower() for part in errors.split(",") if part.strip())
    return RetryPolicy(
        max_attempts=max(1, int(getattr(args, "agent_retry_attempts", 1) or 1)),
        initial_delay_seconds=max(0.0, float(getattr(args, "agent_retry_delay", 0.0) or 0.0)),
        backoff_factor=max(1.0, float(getattr(args, "agent_retry_backoff", 2.0) or 2.0)),
        retryable_errors=retryable_errors or RetryPolicy().retryable_errors,
    )


def _run_review_decision_cli(args):
    decision_path = Path(args.review_decision)
    if not decision_path.exists():
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "error": f"review decision does not exist: {decision_path}",
        }, ensure_ascii=False, indent=2))
        return
    if not args.resume:
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "error": "--resume <manifest-or-run-dir> is required with --review-decision",
        }, ensure_ascii=False, indent=2))
        return
    resume_path = Path(args.resume)
    manifest_path = resume_path / "manifest.json" if resume_path.is_dir() else resume_path
    if not manifest_path.exists():
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "error": f"manifest does not exist: {manifest_path}",
        }, ensure_ascii=False, indent=2))
        return

    decision = json.loads(decision_path.read_text(encoding="utf-8"))
    manager = AgentRunManager(manifest_path.parent, run_id=manifest_path.parent.name)
    result = manager.record_decision(decision)
    result.update({
        "agent_state": "REVIEW_DECISION_RECORDED",
        "state": "REVIEW_DECISION_RECORDED",
    })
    print(json.dumps(result, ensure_ascii=False, indent=2))


def _agent_batch_root(args) -> Path:
    if args.trace:
        trace_path = Path(args.trace)
        if trace_path.exists() and trace_path.is_dir():
            return trace_path
        if trace_path.suffix.lower() == ".jsonl":
            return trace_path.with_suffix("")
        return trace_path
    return PROGRAM_DIR / "agent-runs"


def _run_agent_batch_cli(args):
    catalog_path = Path(args.catalog)
    batch_dir = Path(args.batch)

    if not batch_dir.exists() or not batch_dir.is_dir():
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "error": f"batch directory does not exist: {batch_dir}",
        }, ensure_ascii=False, indent=2))
        return

    if not catalog_path.exists():
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "error": f"catalog does not exist: {catalog_path}",
        }, ensure_ascii=False, indent=2))
        return

    catalog = CatalogLoader(catalog_path).load()
    registry = build_legacy_tool_registry(catalog)
    manager = AgentRunManager(
        _agent_batch_root(args),
        reviews_dir=Path(args.review_output) if args.review_output else None,
    )
    manager.ensure()

    files = sorted([p for p in batch_dir.iterdir() if p.is_file()])
    for file_path in files:
        paths = manager.paths_for(str(file_path))
        trace_store = AgentTraceStore(paths["trace"])
        runner = AgentRunner(registry, trace_store, retry_policy=_agent_retry_policy_from_args(args))
        result = runner.run(run_id=manager.run_id, file_path=str(file_path))
        review_path = str(paths["review"])
        output = _agent_output_from_result(result, paths["trace"], review_path, args.resume or "")
        paths["output"].parent.mkdir(parents=True, exist_ok=True)
        paths["output"].write_text(json.dumps(output, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

        if output["agent_state"] in {AgentState.NEEDS_REVIEW.value, AgentState.FAILED.value}:
            ReviewPackageWriter(paths["review"]).write(build_review_payload(output, str(paths["trace"])))

        manager.record_file(str(file_path), output, paths)

    review_index = manager.write_review_index()
    summary = manager.summary()
    summary.update({
        "agent_state": "BATCH_COMPLETED",
        "state": "BATCH_COMPLETED",
        "review_index": str(review_index),
    })
    print(json.dumps(summary, ensure_ascii=False, indent=2))


def _run_agent_resume(args):
    trace_path = Path(args.resume)
    if trace_path.is_dir() or trace_path.name == "manifest.json":
        _run_agent_batch_resume(args, trace_path)
        return

    if not trace_path.exists():
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "trace": str(trace_path),
            "resume": True,
            "error": f"trace does not exist: {trace_path}",
        }, ensure_ascii=False, indent=2))
        return

    steps = AgentTraceStore(trace_path).load()
    if not steps:
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "trace": str(trace_path),
            "resume": True,
            "error": "trace is empty",
        }, ensure_ascii=False, indent=2))
        return

    last = steps[-1]
    if last.state in {AgentState.COMPLETED, AgentState.NEEDS_REVIEW}:
        output = dict(last.output_summary or {})
        output.update({
            "agent_state": last.state.value,
            "state": last.state.value,
            "trace": str(trace_path),
            "resume": True,
            "file_path": last.file_path,
            "run_id": last.run_id,
        })
        print(json.dumps(output, ensure_ascii=False, indent=2))
        return

    if last.state == AgentState.FAILED:
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "trace": str(trace_path),
            "resume": True,
            "file_path": last.file_path,
            "run_id": last.run_id,
            "error": last.error or "agent run failed",
        }, ensure_ascii=False, indent=2))
        return

    result = _resume_agent_trace(args, trace_path, steps)
    print(json.dumps(result, ensure_ascii=False, indent=2))


def _resume_agent_trace(args, trace_path: Path, steps: list) -> dict:
    last = steps[-1]
    catalog_path = Path(args.catalog)
    if not catalog_path.exists():
        return {
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "trace": str(trace_path),
            "resume": True,
            "file_path": last.file_path,
            "run_id": last.run_id,
            "error": f"catalog does not exist: {catalog_path}",
        }

    catalog = CatalogLoader(catalog_path).load()
    registry = build_legacy_tool_registry(catalog)
    runner = AgentRunner(registry, AgentTraceStore(trace_path), retry_policy=_agent_retry_policy_from_args(args))
    result = runner.resume(run_id=last.run_id, file_path=last.file_path, steps=steps)
    output = _agent_output_from_result(result, trace_path, args.review_output or "", args.resume or "")
    output.update({
        "resume": True,
        "file_path": output.get("file_path") or last.file_path,
        "run_id": last.run_id,
        "last_state": last.state.value,
    })
    if output["agent_state"] in {AgentState.NEEDS_REVIEW.value, AgentState.FAILED.value} and args.review_output:
        ReviewPackageWriter(Path(args.review_output)).write(build_review_payload(output, str(trace_path)))
    return output


def _run_agent_batch_resume(args, resume_path: Path):
    manifest_path = resume_path / "manifest.json" if resume_path.is_dir() else resume_path
    if not manifest_path.exists():
        print(json.dumps({
            "agent_state": AgentState.FAILED.value,
            "state": AgentState.FAILED.value,
            "resume": True,
            "error": f"manifest does not exist: {manifest_path}",
        }, ensure_ascii=False, indent=2))
        return

    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    files = manifest.get("files", [])
    run_dir = manifest_path.parent
    review_dir = None
    for item in files:
        if item.get("review"):
            review_dir = Path(item["review"]).parent
            break
    manager = AgentRunManager(run_dir, run_id=run_dir.name, reviews_dir=review_dir)
    manager.ensure()
    resumed = []
    skipped = []
    for item in files:
        state = item.get("agent_state", "")
        if state in {AgentState.COMPLETED.value, AgentState.NEEDS_REVIEW.value, AgentState.FAILED.value}:
            skipped.append(item)
            continue
        trace_path = Path(item.get("trace", ""))
        steps = AgentTraceStore(trace_path).load()
        if not steps:
            resumed.append({
                "file_path": item.get("file_path", ""),
                "agent_state": AgentState.FAILED.value,
                "error": "trace is empty or missing",
            })
            continue
        output = _resume_agent_trace(args, trace_path, steps)
        paths = {
            "file_id": item.get("file_id") or Path(item.get("output", trace_path.stem)).stem,
            "trace": trace_path,
            "review": Path(item.get("review") or manager.paths_for(item.get("file_path", ""))["review"]),
            "output": Path(item.get("output") or manager.paths_for(item.get("file_path", ""))["output"]),
        }
        paths["output"].parent.mkdir(parents=True, exist_ok=True)
        paths["output"].write_text(json.dumps(output, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        if output["agent_state"] in {AgentState.NEEDS_REVIEW.value, AgentState.FAILED.value}:
            ReviewPackageWriter(paths["review"]).write(build_review_payload(output, str(trace_path)))
        manager.record_file(item.get("file_path", output.get("file_path", "")), output, paths)
        resumed.append(output)

    review_index = manager.write_review_index()
    print(json.dumps({
        "agent_state": "BATCH_RESUMED",
        "state": "BATCH_RESUMED",
        "resume": True,
        "manifest": str(manifest_path),
        "review_index": str(review_index),
        "resumed_count": len(resumed),
        "skipped_count": len(skipped),
        "resumed": resumed,
        "skipped": skipped,
    }, ensure_ascii=False, indent=2))


def main():
    parser = argparse.ArgumentParser(
        description="基于 VLM+LLM 的文件目录推测系统",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("file", nargs="?", help="待推测的文件路径")
    parser.add_argument("--batch", help="批量处理目录下所有文件")
    parser.add_argument("--catalog", default=str(DEFAULT_CATALOG),
                        help=f"编目规则 xlsx 路径（默认 {DEFAULT_CATALOG}）")
    parser.add_argument("--json", action="store_true", help="以 JSON 输出结果")
    parser.add_argument("--output", help="运行结果保存路径（默认程序目录下 file_directory_predictor_output.txt）")
    parser.add_argument("--log", help="运行日志保存路径（默认程序目录下 file_directory_predictor.log）")
    parser.add_argument("--agent", action="store_true", help="启用智能体状态机执行")
    parser.add_argument("--trace", help="智能体执行轨迹 JSONL 保存路径")
    parser.add_argument("--review-output", help="需要人工复核时输出复核材料的路径")
    parser.add_argument("--resume", help="从已有 trace JSONL 恢复智能体任务")
    parser.add_argument("--agent-retry-attempts", type=int, default=1, help="智能体工具调用最大尝试次数")
    parser.add_argument("--agent-retry-delay", type=float, default=0.0, help="智能体重试初始等待秒数")
    parser.add_argument("--agent-retry-backoff", type=float, default=2.0, help="智能体重试退避倍率")
    parser.add_argument("--agent-retry-errors", default="", help="逗号分隔的可重试错误关键字")
    parser.add_argument("--agent-preflight", action="store_true", help="检查智能体模型配置，不调用网络 API")
    parser.add_argument("--review-decision", help="写入人工复核决定 JSON，并更新 run manifest")
    args = parser.parse_args()

    if not args.file and not args.batch and not (args.agent and (args.resume or args.agent_preflight or args.review_decision)):
        parser.error("请提供文件路径或 --batch 目录")

    output_path = Path(args.output) if args.output else PROGRAM_DIR / DEFAULT_OUTPUT_FILE
    log_path = Path(args.log) if args.log else PROGRAM_DIR / DEFAULT_LOG_FILE

    with _save_cli_streams(output_path, log_path):
        if args.agent:
            _run_agent_cli(args)
            return

        catalog_path = Path(args.catalog)
        if not catalog_path.exists():
            print(f"[错误] 编目规则文件不存在: {catalog_path}", file=sys.stderr)
            sys.exit(1)

        print(f"[加载编目规则] {catalog_path}", file=sys.stderr)
        catalog = CatalogLoader(catalog_path).load()
        print(f"  共 {len(catalog.entries)} 条规则，{len(catalog.case_types)} 种案件类型，"
              f"{len(catalog.categories)} 种材料类别", file=sys.stderr)

        predictor = DirectoryPredictor(catalog)

        if args.batch:
            batch_dir = Path(args.batch)
            results = []
            files = sorted([p for p in batch_dir.iterdir() if p.is_file()])
            for f in files:
                print(f"\n[处理] {f.name}", file=sys.stderr)
                try:
                    r = predictor.predict(f)
                    results.append(r.to_dict())
                    _print_result(r, as_json=args.json)
                except Exception as e:
                    print(f"  [失败] {e}", file=sys.stderr)
            # 批量结果汇总
            if args.json:
                print(json.dumps(results, ensure_ascii=False, indent=2))
            return

        file_path = Path(args.file)
        if not file_path.exists():
            print(f"[错误] 文件不存在: {file_path}", file=sys.stderr)
            sys.exit(1)
        result = predictor.predict(file_path)
        _print_result(result, as_json=args.json)


if __name__ == "__main__":
    main()
